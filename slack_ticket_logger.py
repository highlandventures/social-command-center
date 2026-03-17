#!/usr/bin/env python3
"""
Slack Ticket Logger
-------------------
A lightweight Flask server that receives webhook payloads (from Zapier)
and appends formatted ticket entries to a single running .docx file.

Endpoints:
    POST /ticket  — Append a new ticket entry
    GET  /health  — Health check

Usage:
    pip install flask python-docx
    python slack_ticket_logger.py
"""

import os
import json
from datetime import datetime
from pathlib import Path

from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# ── Configuration ───────────────────────────────────────────────────────────
DOCX_PATH = os.path.expanduser("~/social command/slack_tickets.docx")
PORT = 5050


def ensure_docx_exists():
    """Create the .docx with a title page if it doesn't exist yet."""
    path = Path(DOCX_PATH)
    if path.exists():
        return

    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(6)
    run = title.add_run("Slack Ticket Log")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(24)
    run = subtitle.add_run("Auto-generated from Slack submissions via Zapier webhook")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Divider
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("— — —")
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    doc.save(str(path))
    print(f"[INIT] Created new ticket log at {DOCX_PATH}")


def format_timestamp():
    """Return a human-readable timestamp, e.g. 'March 16, 2026 at 2:30 PM'."""
    now = datetime.now()
    return now.strftime("%B %-d, %Y at %-I:%M %p")


def append_ticket(fields: dict):
    """Open the .docx, append a formatted ticket entry, and save."""
    ensure_docx_exists()
    doc = Document(DOCX_PATH)

    # ── Visual divider ──────────────────────────────────────────────────
    divider = doc.add_paragraph()
    divider.space_before = Pt(18)
    divider.space_after = Pt(4)
    run = divider.add_run("━" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ── Ticket heading with timestamp ───────────────────────────────────
    heading = doc.add_paragraph()
    heading.space_after = Pt(8)
    run = heading.add_run(f"Ticket — {format_timestamp()}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── Dynamic field entries ───────────────────────────────────────────
    if not fields:
        p = doc.add_paragraph()
        run = p.add_run("(No fields submitted)")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        for key, value in fields.items():
            p = doc.add_paragraph()
            p.space_after = Pt(2)

            # Bold label
            label_run = p.add_run(f"{key}: ")
            label_run.bold = True
            label_run.font.size = Pt(11)

            # Value
            value_run = p.add_run(str(value))
            value_run.font.size = Pt(11)

    doc.save(DOCX_PATH)


# ── Routes ──────────────────────────────────────────────────────────────────

@app.route("/ticket", methods=["POST"])
def receive_ticket():
    """Accept a ticket payload (JSON or form-encoded) and append to .docx."""
    # Parse payload from either content type
    if request.is_json:
        fields = request.get_json(force=True)
    else:
        fields = request.form.to_dict()

    # Also try to parse a raw body as JSON if fields are empty
    if not fields and request.data:
        try:
            fields = json.loads(request.data)
        except (json.JSONDecodeError, TypeError):
            fields = {}

    append_ticket(fields)

    timestamp = format_timestamp()
    field_count = len(fields)
    print(f"[TICKET] Appended {field_count} field(s) at {timestamp}")

    return jsonify({
        "status": "ok",
        "message": f"Ticket logged with {field_count} field(s)",
        "timestamp": timestamp
    }), 200


@app.route("/health", methods=["GET"])
def health():
    """Health check endpoint."""
    docx_exists = os.path.exists(DOCX_PATH)
    return jsonify({
        "status": "healthy",
        "docx_path": DOCX_PATH,
        "docx_exists": docx_exists,
        "port": PORT
    }), 200


# ── Main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    ensure_docx_exists()
    print(f"Slack Ticket Logger running on http://localhost:{PORT}")
    print(f"  POST /ticket  — Append a ticket")
    print(f"  GET  /health  — Health check")
    print(f"  Saving to: {DOCX_PATH}")
    print()
    app.run(host="0.0.0.0", port=PORT, debug=False)
